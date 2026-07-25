"""Document parsing: PDF, DOCX, TXT, and images -> per-page plain text.

Every chunk downstream carries a ``page`` number, so parsing must preserve
page boundaries even for formats without a native page concept (DOCX, TXT,
and images are treated as a single page). Scanned PDF pages and image
uploads (PNG/JPG/WEBP) with no text layer go through pytesseract OCR.
"""

import io
from dataclasses import dataclass

import docx
import fitz  # pymupdf
import pytesseract
from PIL import Image

from app.constants.upload import ALLOWED_EXTENSIONS, IMAGE_EXTENSIONS
from app.utils.errors import create_error

# Render scale for OCR-ing scanned PDF pages: 2x default 72dpi puts small
# print at a resolution tesseract can actually read.
_OCR_RENDER_MATRIX = fitz.Matrix(2, 2)

# Tesseract's whole-page automatic layout segmentation can silently drop an
# isolated mid-page table block on a sparse, multi-table layout (verified:
# a bill-of-lading's commodity/weight table vanished entirely from a
# whole-page OCR pass, but read back perfectly once cropped tightly to its
# own region). Re-OCR-ing in narrow horizontal bands sidesteps this — a
# high whitespace-to-content ratio is what confuses the segmenter, and thin
# bands keep that ratio low even for an isolated table row. Bands overlap so
# a text row landing on a seam still gets read cleanly in the neighboring
# band; the resulting duplication/occasional garbled seam text is an
# acceptable tradeoff on this fallback path — a chunk having the fact once,
# cleanly, is what matters for retrieval.
_OCR_BAND_COUNT = 16
_OCR_BAND_OVERLAP_PX = 40


@dataclass
class PageText:
    """Plain text extracted from one page of a source document."""

    page: int
    text: str


def parse_document(*, filename: str, extension: str, content: bytes) -> list[PageText]:
    """Dispatch to the right parser by file extension. Raises AppError (415)
    for unsupported extensions."""
    if extension == "pdf":
        return _parse_pdf(content)
    if extension == "docx":
        return _parse_docx(content)
    if extension == "txt":
        return _parse_txt(content)
    if extension in IMAGE_EXTENSIONS:
        return _parse_image(content)
    raise create_error(
        message="Unsupported file type",
        why=f"'{extension}' is not one of {sorted(ALLOWED_EXTENSIONS)}",
        fix="Upload a PDF, DOCX, TXT, PNG, JPG, or WEBP file",
        status_code=415,
        filename=filename,
    )


def _ocr_in_bands(image: Image.Image) -> str:
    """OCR horizontal strips of the page independently and concatenate. Each
    strip overlaps its neighbor slightly so a text row isn't cleanly split by
    a band seam."""
    width, height = image.size
    band_height = height // _OCR_BAND_COUNT
    parts: list[str] = []
    for index in range(_OCR_BAND_COUNT):
        top = max(0, index * band_height - _OCR_BAND_OVERLAP_PX)
        bottom = height if index == _OCR_BAND_COUNT - 1 else (index + 1) * band_height
        band = image.crop((0, top, width, bottom))
        text = pytesseract.image_to_string(band).strip()
        if text:
            parts.append(text)
    return "\n".join(parts).strip()


def _ocr_image(image: Image.Image) -> str:
    """OCR a full page, keeping whichever of (whole-page pass, banded pass)
    extracted more text — the cheap whole-page pass is usually sufficient,
    but banding recovers content it silently drops (see module docstring)."""
    whole_page_text = pytesseract.image_to_string(image).strip()
    banded_text = _ocr_in_bands(image)
    return banded_text if len(banded_text) > len(whole_page_text) else whole_page_text


def _parse_pdf(content: bytes) -> list[PageText]:
    pages: list[PageText] = []
    with fitz.open(stream=content, filetype="pdf") as document:
        for index, page in enumerate(document):
            text = page.get_text().strip()
            if not text:
                # Scanned/image-only page: no text layer, so fall back to OCR
                # on a rendered bitmap of the page.
                pixmap = page.get_pixmap(matrix=_OCR_RENDER_MATRIX)
                image = Image.open(io.BytesIO(pixmap.tobytes("png")))
                text = _ocr_image(image)
            if text:
                pages.append(PageText(page=index + 1, text=text))
    return pages


def _parse_docx(content: bytes) -> list[PageText]:
    """DOCX has no native page concept, so the whole document is page 1.
    Tables are row-joined ("Cell A | Cell B") and appended after the running
    paragraph text — rate confirmations are table-heavy and paragraph-only
    extraction silently drops the shipment fields that live in table cells.
    """
    document = docx.Document(io.BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    return [PageText(page=1, text=text)] if text else []


def _parse_txt(content: bytes) -> list[PageText]:
    text = content.decode("utf-8", errors="replace").strip()
    return [PageText(page=1, text=text)] if text else []


def _parse_image(content: bytes) -> list[PageText]:
    image = Image.open(io.BytesIO(content))
    text = _ocr_image(image)
    return [PageText(page=1, text=text)] if text else []
